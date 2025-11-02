from typing import List, Dict, Any
from utils import format_ts
import logging
import os


def export_txt(path: str, segments: List[Dict[str, Any]], speaker_map: Dict[int, str], include_speakers: bool = True):
    logging.info("[FileSystem] Exporting TXT to: %s", path)
    logging.info("[FileSystem] Output directory: %s", os.path.dirname(path))
    logging.info("[FileSystem] Output directory exists: %s", os.path.exists(os.path.dirname(path)))

    lines = []
    for seg in segments:
        timestamp = f"[{format_ts(seg['start'])} – {format_ts(seg['end'])}]"
        if include_speakers:
            spk_name = speaker_map.get(seg.get("speaker", 0), f"Speaker {seg.get('speaker', 0)+1}")
            lines.append(f"{timestamp} {spk_name}: {seg['text']}")
        else:
            lines.append(f"{timestamp} {seg['text']}")
    content = "\n".join(lines)

    try:
        logging.info("[FileSystem] Writing %d bytes to TXT file", len(content.encode('utf-8')))
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content + "\n")
        logging.info("[FileSystem] TXT file written successfully")
        logging.info("[FileSystem] TXT file size: %d bytes", os.path.getsize(path))
    except Exception as e:
        logging.error("[FileSystem] Failed to write TXT file: %s", str(e), exc_info=True)
        raise


def export_docx(path: str, segments: List[Dict[str, Any]], speaker_map: Dict[int, str], include_speakers: bool = True):
    logging.info("[FileSystem] Exporting DOCX to: %s", path)
    logging.info("[FileSystem] Output directory: %s", os.path.dirname(path))
    logging.info("[FileSystem] Output directory exists: %s", os.path.exists(os.path.dirname(path)))

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()
        doc.add_heading('Transcript', level=1)

        for seg in segments:
            p = doc.add_paragraph()
            run1 = p.add_run(f"[{format_ts(seg['start'])} – {format_ts(seg['end'])}] ")
            if include_speakers:
                spk_name = speaker_map.get(seg.get("speaker", 0), f"Speaker {seg.get('speaker', 0)+1}")
                run2 = p.add_run(f"{spk_name}: ")
                run2.bold = True
            p.add_run(seg['text'])

        logging.info("[FileSystem] Saving DOCX document")
        doc.save(path)
        logging.info("[FileSystem] DOCX file written successfully")
        logging.info("[FileSystem] DOCX file size: %d bytes", os.path.getsize(path))
    except Exception as e:
        logging.error("[FileSystem] Failed to write DOCX file: %s", str(e), exc_info=True)
        raise
